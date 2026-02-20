"""
Generate a professional Word document (.docx) summarising the new features
implemented on the IAG DENUCHANGE Workshop 2026 website.

Images are left as numbered placeholders — the author inserts them manually.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT = Path(__file__).parent / "DENUCHANGE_Website_Report.docx"

# ── colours ──────────────────────────────────────────────────────────────
TEAL = RGBColor(0x0A, 0x6E, 0x84)      # primary brand colour
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_colour: str):
    """Apply a solid background fill to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_colour,
    })
    shading.append(shd)


def add_heading_styled(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL
    return h


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p


def add_image_placeholder(doc: Document, number: int, caption: str):
    """Insert a light-grey box with '[Εικόνα X]' and a caption underneath."""
    # Placeholder box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "F3F4F6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n[Εικόνα {number}]\n\n")
    run.font.size = Pt(14)
    run.font.color.rgb = GREY
    run.font.italic = True
    # set cell width
    cell.width = Cm(14)

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Εικόνα {number}: {caption}")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p


# ── build document ──────────────────────────────────────────────────────

def build():
    doc = Document()

    # ------ page margins ------
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ------ default font ------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK

    # ══════════════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("IAG DENUCHANGE Workshop 2026")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = TEAL

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    r = subtitle.add_run("Αναφορά Υλοποίησης Ιστοσελίδας")
    r.font.size = Pt(14)
    r.font.color.rgb = GREY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(20)
    r = date_p.add_run("Φεβρουάριος 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

    # ══════════════════════════════════════════════════════════════════════
    # INTRO
    # ══════════════════════════════════════════════════════════════════════
    add_body(doc,
        "Το παρόν έγγραφο παρουσιάζει αναλυτικά τις νέες λειτουργίες που υλοποιήθηκαν "
        "στην ιστοσελίδα του Workshop και δεν είναι άμεσα ορατές κατά την απλή πλοήγηση: "
        "φόρμες εγγραφής και υποβολών, αυτόματα emails επιβεβαίωσης, έλεγχοι διπλοεγγραφής, "
        "επαλήθευση email, και σύνδεση με βάση δεδομένων."
    )
    add_body(doc,
        "Σε κάθε ενότητα υπάρχουν κενά πλαίσια εικόνων (placeholders) με αριθμούς και "
        "λεζάντες. Οι εικόνες (screenshots) θα εισαχθούν χειροκίνητα."
    )

    img_num = 0  # running counter for image placeholders

    # ══════════════════════════════════════════════════════════════════════
    # 1. REGISTRATION FORM
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Φόρμα Εγγραφής (Registration)", level=1)

    add_body(doc,
        "Η φόρμα εγγραφής αντικατέστησε την παλαιά φόρμα Google. Ο χρήστης συμπληρώνει "
        "τα στοιχεία του σε ένα modal dialog απευθείας στη σελίδα, χωρίς ανακατεύθυνση."
    )

    add_body(doc, "Πεδία φόρμας:")
    add_bullet(doc, "Όνομα, Επώνυμο, Email, Affiliation, Χώρα")
    add_bullet(doc, "Τύπος εγγραφής: Regular – Full (€450), Student – Full (€300), Meeting Only (€100), Accompanying Person (€300)")
    add_bullet(doc, "Πρόθεση παρουσίασης: Oral / Poster / Καμία")
    add_bullet(doc, "Διατροφικές ανάγκες (Vegetarian, Vegan, Kosher, Gluten-free, Other)")
    add_bullet(doc, "Ειδικές απαιτήσεις (προαιρετικό)")

    img_num += 1
    add_image_placeholder(doc, img_num, "Η φόρμα εγγραφής ανοιχτή στη σελίδα (Registration Dialog).")

    # ── email verification ──
    add_heading_styled(doc, "1.1 Επαλήθευση Email (Verification)", level=2)

    add_body(doc,
        "Πριν την υποβολή, ο χρήστης πατάει «Verify Email». Αποστέλλεται αυτόματα ένα "
        "email επαλήθευσης στη διεύθυνσή του. Μόλις επιβεβαιώσει ότι το έλαβε, μπορεί "
        "να ολοκληρώσει τη φόρμα. Ο μηχανισμός αυτός αποτρέπει τυπογραφικά λάθη στο email."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Κουμπί «Verify Email» και κατάσταση επαλήθευσης στη φόρμα.")

    img_num += 1
    add_image_placeholder(doc, img_num, "Το email επαλήθευσης που λαμβάνει ο χρήστης (Email Verification).")

    # ── duplicate check ──
    add_heading_styled(doc, "1.2 Έλεγχος Διπλοεγγραφής", level=2)

    add_body(doc,
        "Αν ο χρήστης προσπαθήσει να εγγραφεί δεύτερη φορά με το ίδιο email, δεν "
        "δημιουργείται νέα εγγραφή. Εμφανίζεται μήνυμα που τον ενημερώνει ότι η "
        "εγγραφή υπάρχει ήδη, με σύνδεσμο επικοινωνίας (mailto link) για τυχόν αλλαγές. "
        "Δεν αποκαλύπτονται προσωπικά στοιχεία."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Μήνυμα διπλοεγγραφής με σύνδεσμο «contact us».")

    # ── confirmation email ──
    add_heading_styled(doc, "1.3 Email Επιβεβαίωσης Εγγραφής", level=2)

    add_body(doc,
        "Μετά την επιτυχή εγγραφή, αποστέλλεται αυτόματα email επιβεβαίωσης στον "
        "συμμετέχοντα. Το email περιλαμβάνει:"
    )
    add_bullet(doc, "Banner με τον τίτλο του Workshop")
    add_bullet(doc, "Πίνακα με τα στοιχεία εγγραφής (όνομα, email, affiliation, τύπος, κ.λπ.)")
    add_bullet(doc, "Οδηγίες για το επόμενο βήμα (πληρωμή)")
    add_bullet(doc, "Footer με σύνδεσμο επικοινωνίας για αλλαγές")

    add_body(doc,
        "Τα emails αποστέλλονται από τη διεύθυνση denuchange.workshop.2026@gmail.com."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Email επιβεβαίωσης εγγραφής (Registration Confirmed).")

    # ── success state ──
    add_heading_styled(doc, "1.4 Μήνυμα Επιτυχίας", level=2)

    add_body(doc,
        "Μετά την υποβολή, η φόρμα εμφανίζει μήνυμα επιτυχίας με εικονίδιο ✓ που "
        "ενημερώνει τον χρήστη ότι η εγγραφή ολοκληρώθηκε και ότι θα λάβει email."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Μήνυμα επιτυχίας μετά την εγγραφή (Success State).")

    # ══════════════════════════════════════════════════════════════════════
    # 2. ABSTRACT SUBMISSION FORM
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Φόρμα Υποβολής Περίληψης (Abstract Submission)", level=1)

    add_body(doc,
        "Ξεχωριστή φόρμα αποκλειστικά για την υποβολή abstract. Ο χρήστης εισάγει "
        "πρώτα τη διεύθυνση email που χρησιμοποίησε κατά την εγγραφή. Η φόρμα ελέγχει "
        "αυτόματα αν υπάρχει εγγραφή, ανακτά τα στοιχεία (όνομα, affiliation) και τα "
        "συμπληρώνει αυτόματα."
    )

    add_body(doc, "Πεδία φόρμας:")
    add_bullet(doc, "Email (lookup εγγραφής)")
    add_bullet(doc, "Τίτλος abstract")
    add_bullet(doc, "Συν-συγγραφείς (προαιρετικό)")
    add_bullet(doc, "Τύπος παρουσίασης: Oral / Poster")
    add_bullet(doc, "Αρχείο abstract (.docx, μέχρι 5 MB)")

    img_num += 1
    add_image_placeholder(doc, img_num, "Η φόρμα υποβολής abstract (Abstract Submission Dialog).")

    # ── registration lookup ──
    add_heading_styled(doc, "2.1 Αυτόματος Έλεγχος Εγγραφής", level=2)

    add_body(doc,
        "Μόλις ο χρήστης εισάγει email, η φόρμα ελέγχει αν υπάρχει αντίστοιχη εγγραφή "
        "στη βάση. Αν δεν βρεθεί, εμφανίζεται μήνυμα «This email is not registered» και "
        "η υποβολή δεν επιτρέπεται. Ο έλεγχος γίνεται case-insensitive."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Μήνυμα «Email not registered» στη φόρμα abstract.")

    # ── duplicate check ──
    add_heading_styled(doc, "2.2 Έλεγχος Διπλής Υποβολής", level=2)

    add_body(doc,
        "Αν ο χρήστης έχει ήδη υποβάλει abstract με το ίδιο email, εμφανίζεται "
        "αντίστοιχο μήνυμα με σύνδεσμο επικοινωνίας."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Μήνυμα «abstract already submitted» με σύνδεσμο «contact us».")

    # ── email ──
    add_heading_styled(doc, "2.3 Email Επιβεβαίωσης Υποβολής", level=2)

    add_body(doc,
        "Μετά την επιτυχή υποβολή, αποστέλλεται αυτόματα email με τα στοιχεία "
        "υποβολής: τίτλο abstract, τύπο παρουσίασης, συν-συγγραφείς."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Email επιβεβαίωσης υποβολής abstract (Abstract Received).")

    # ══════════════════════════════════════════════════════════════════════
    # 3. PAYMENT RECEIPT UPLOAD
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Ανάρτηση Αποδεικτικού Πληρωμής (Payment Receipt)", level=1)

    add_body(doc,
        "Ο χρήστης μπορεί να ανεβάσει σκαναρισμένο αποδεικτικό πληρωμής (PDF/εικόνα) "
        "μέσω ειδικής φόρμας. Η ταυτοποίηση γίνεται μέσω του email εγγραφής."
    )

    add_body(doc, "Πεδία φόρμας:")
    add_bullet(doc, "Email (lookup εγγραφής)")
    add_bullet(doc, "Αρχείο αποδεικτικού (PDF, JPG, PNG — μέχρι 5 MB)")
    add_bullet(doc, "Σημειώσεις (προαιρετικό)")

    img_num += 1
    add_image_placeholder(doc, img_num, "Η φόρμα ανάρτησης αποδεικτικού πληρωμής (Payment Receipt Dialog).")

    # ── duplicate check ──
    add_heading_styled(doc, "3.1 Έλεγχος Διπλής Ανάρτησης", level=2)

    add_body(doc,
        "Αν έχει ήδη ανέβει αποδεικτικό για το ίδιο email, εμφανίζεται μήνυμα "
        "με σύνδεσμο επικοινωνίας."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Μήνυμα «payment receipt already uploaded» με σύνδεσμο «contact us».")

    # ── email ──
    add_heading_styled(doc, "3.2 Email Επιβεβαίωσης Πληρωμής", level=2)

    add_body(doc,
        "Μετά την ανάρτηση, αποστέλλεται email που επιβεβαιώνει τη λήψη του "
        "αποδεικτικού."
    )

    img_num += 1
    add_image_placeholder(doc, img_num, "Email επιβεβαίωσης λήψης αποδεικτικού (Payment Receipt Received).")

    # ══════════════════════════════════════════════════════════════════════
    # 4. DATABASE
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Βάση Δεδομένων (Supabase)", level=1)

    add_body(doc,
        "Όλα τα δεδομένα αποθηκεύονται σε ασφαλή βάση δεδομένων Supabase (PostgreSQL) "
        "με ενεργοποιημένο Row-Level Security (RLS). "
        "Τα αρχεία (abstracts, αποδεικτικά) αποθηκεύονται σε Storage buckets."
    )

    add_body(doc, "Πίνακες:")
    add_bullet(doc, "registrations — στοιχεία εγγραφής κάθε συμμετέχοντα")
    add_bullet(doc, "abstracts — μεταδεδομένα abstract, link στο αρχείο")
    add_bullet(doc, "payment_receipts — μεταδεδομένα πληρωμής, link στο αρχείο")

    add_body(doc, "Storage buckets:")
    add_bullet(doc, "abstracts — αρχεία .docx περιλήψεων")
    add_bullet(doc, "payment-receipts — αρχεία αποδεικτικών (PDF/εικόνα)")

    img_num += 1
    add_image_placeholder(doc, img_num,
        "Πίνακας registrations στο Supabase Dashboard — λίστα εγγεγραμμένων.")

    img_num += 1
    add_image_placeholder(doc, img_num,
        "Πίνακας abstracts στο Supabase Dashboard — υποβληθείσες περιλήψεις.")

    img_num += 1
    add_image_placeholder(doc, img_num,
        "Πίνακας payment_receipts στο Supabase Dashboard — αποδεικτικά πληρωμής.")

    img_num += 1
    add_image_placeholder(doc, img_num,
        "Storage bucket «abstracts» — αποθηκευμένα αρχεία περιλήψεων.")

    img_num += 1
    add_image_placeholder(doc, img_num,
        "Storage bucket «payment-receipts» — αποθηκευμένα αποδεικτικά.")

    # ══════════════════════════════════════════════════════════════════════
    # 5. OFFICIAL EMAIL
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Επίσημο Email Συνεδρίου", level=1)

    add_body(doc,
        "Δημιουργήθηκε η διεύθυνση denuchange.workshop.2026@gmail.com αποκλειστικά "
        "για το Workshop. Μέσω αυτής αποστέλλονται αυτόματα όλα τα emails επιβεβαίωσης "
        "και λαμβάνονται τα αιτήματα αλλαγών."
    )

    img_num += 1
    add_image_placeholder(doc, img_num,
        "Inbox του denuchange.workshop.2026@gmail.com — αυτόματα emails επιβεβαίωσης.")

    # ══════════════════════════════════════════════════════════════════════
    # 6. PENDING
    # ══════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Εκκρεμότητες", level=1)

    add_body(doc,
        "Τα παρακάτω αναμένουν απόφαση από την ομάδα:"
    )
    add_bullet(doc,
        "Στοιχεία τραπέζης (IBAN): Μόλις οριστικοποιηθούν, θα προστεθούν στη σελίδα "
        "και στα emails.)."
    )
    add_bullet(doc,
        "Κατηγορία εκδρομής: Να αποφασιστεί αν θα υπάρξει ξεχωριστή κατηγορία εγγραφής "
        "μόνο για τη συμμετοχή στην εκδρομή, με αντίστοιχο ποσό."
    )

    # ══════════════════════════════════════════════════════════════════════
    # CLOSING
    # ══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()  # spacer
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(16)
    r = closing.add_run("Με εκτίμηση,")
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    name_p = doc.add_paragraph()
    r = name_p.add_run("Αλέξανδρος Λιάσκος")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = DARK

    # ── save ──
    doc.save(str(OUTPUT))
    print(f"✓ Document saved to {OUTPUT}")
    print(f"  Total image placeholders: {img_num}")


if __name__ == "__main__":
    build()
